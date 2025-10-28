import streamlit as st
import google.generativeai as genai
import pandas as pd
import json
import requests
import time
from datetime import datetime
import PyPDF2
import io

# DOCX support for Word export
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("⚠️ python-docx not installed. Install: pip install python-docx")

# App config
st.set_page_config(page_title="Qforia Research Platform", layout="wide")
st.title("fmc Content Guide")

# Grok API Configuration
GROK_API_URL = "https://api.x.ai/v1/chat/completions"

# Initialize session states
if 'fanout_results' not in st.session_state:
    st.session_state.fanout_results = None
if 'generation_details' not in st.session_state:
    st.session_state.generation_details = None
if 'research_results' not in st.session_state:
    st.session_state.research_results = {}
if 'selected_queries' not in st.session_state:
    st.session_state.selected_queries = set()
if 'pdf_analysis' not in st.session_state:
    st.session_state.pdf_analysis = None
if 'enhanced_topics' not in st.session_state:
    st.session_state.enhanced_topics = []
if 'generated_content' not in st.session_state:
    st.session_state.generated_content = {}
if 'content_structure' not in st.session_state:
    st.session_state.content_structure = []
if 'user_query' not in st.session_state:
    st.session_state.user_query = ''
if 'user_intent' not in st.session_state:
    st.session_state.user_intent = ''

# Sidebar Configuration
st.sidebar.header("🔧 Configuration")
gemini_key = st.sidebar.text_input("Gemini API Key", type="password")
perplexity_key = st.sidebar.text_input("Perplexity API Key", type="password")
grok_key = st.sidebar.text_input("Grok API Key (for content generation)", type="password", help="Enter your xAI Grok API key")

# Gemini model selector with WORKING models (no models/ prefix needed)
st.sidebar.subheader("Gemini Model")
gemini_model = st.sidebar.selectbox(
    "Select Model",
    [
        "gemini-2.0-flash-exp",
        "gemini-2.5-flash",
        "gemini-2.5-pro",
        "gemini-2.0-flash",
    ],
    index=0
)

# Configure Gemini with working model
if gemini_key:
    genai.configure(api_key=gemini_key)
    model = genai.GenerativeModel(gemini_model)
    st.sidebar.success(f"✅ {gemini_model}")
else:
    model = None

# Utility Functions
def call_perplexity(query, system_prompt="Provide comprehensive, actionable insights with specific data points, statistics, and practical recommendations."):
    if not perplexity_key:
        return {"error": "Missing Perplexity API key"}
    
    headers = {
        "Authorization": f"Bearer {perplexity_key}",
        "Content-Type": "application/json",
        "Accept": "application/json"
    }
    data = {
        "model": "sonar",  # CHEAPEST model - $0.2 per 1M tokens (vs sonar-pro at $3)
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": query}
        ],
        "max_tokens": 800  # Increased from 500 for more comprehensive answers
    }
    
    try:
        response = requests.post(
            "https://api.perplexity.ai/chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": f"Perplexity API error: {e}"}

def call_grok(messages, max_tokens=4000, temperature=0.7):
    """Call Grok API for content generation"""
    if not grok_key:
        return None, "Missing Grok API key"
    
    headers = {
        "Authorization": f"Bearer {grok_key}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "messages": messages,
        "model": "grok-2-1212",
        "stream": False,
        "temperature": temperature,
        "max_tokens": max_tokens
    }
    
    try:
        response = requests.post(GROK_API_URL, headers=headers, json=payload, timeout=120)
        response.raise_for_status()
        result = response.json()
        
        if 'choices' in result and len(result['choices']) > 0:
            return result['choices'][0]['message']['content'], None
        return None, "Unexpected response format"
    except Exception as e:
        return None, f"Grok API error: {str(e)}"

def generate_content_structure_from_queries(research_data, topic, selected_queries):
    """Generate article structure directly from selected query fan-out queries"""
    if not grok_key:
        return None, "Grok API key required"
    
    # Get the actual queries from selected items
    sections_list = []
    for query_id in selected_queries:
        if query_id in research_data:
            query_text = research_data[query_id]['query']
            sections_list.append({
                "title": query_text,  # FIXED: Using exact query text as section heading
                "query_id": query_id,
                "description": f"Content based on: {query_text}",
                "key_points": [],
                "needs_table": True,  # Enable tables by default
                "table_description": f"Data and comparisons for {query_text}",
                "needs_infographic": False,
                "infographic_description": "",
                "estimated_words": 600
            })
    
    if not sections_list:
        return None, "No queries selected"
    
    # Create structure
    structure = {
        "article_title": topic,
        "meta_description": f"Comprehensive guide about {topic} with detailed analysis and insights.",
        "primary_keyword": topic,
        "semantic_keywords": [topic],
        "sections": sections_list
    }
    
    return structure, None

def generate_section_content(section, research_context, topic, user_intent):
    """Generate data-driven content with emphasis on tables, bullet points, and structured data"""
    if not grok_key:
        return None, "Grok API key required"
    
    # Get research data for this section
    section_research = ""
    if 'query_id' in section and section['query_id'] in research_context:
        research_data = research_context[section['query_id']]
        if 'result' in research_data:
            result = research_data['result']
            if isinstance(result, dict) and 'choices' in result:
                section_research = result['choices'][0]['message']['content']
            elif isinstance(result, str):
                section_research = result
    
    # ENHANCED: Stricter check for insufficient data
    if not section_research or len(section_research.strip()) < 100:
        return None, "INSUFFICIENT_DATA: Not enough research content available"
    
    # Check for actual useful content (not just error messages)
    if "error" in section_research.lower() or "not available" in section_research.lower():
        return None, "INSUFFICIENT_DATA: Research data contains errors or unavailable information"
    
    # ENHANCED: Aggressive data-focused prompt with user intent integration
    prompt = f"""Create EXTREMELY data-dense, structured content optimized for maximum information value.

ARTICLE TOPIC: {topic}
SECTION TITLE: {section['title']}

USER INTENT/PURPOSE (CRITICAL - Optimize output for this):
{user_intent}

RESEARCH DATA:
{section_research}

ABSOLUTE REQUIREMENTS - NO EXCEPTIONS:
1. ⚡ DATA FIRST: 80% structured data (bullets, lists, tables), 20% prose maximum
2. 📊 EXTRACT EVERY NUMBER: All statistics, prices, percentages, dates, metrics, measurements
3. 🎯 BULLET POINTS ONLY: Use bullets and numbered lists - NO paragraph blocks
4. ❌ INSUFFICIENT DATA RESPONSE: If research lacks concrete data/numbers, return EXACTLY: "INSUFFICIENT_DATA"
5. 📈 COMPARATIVE DATA: Show comparisons, before/after, trends, rankings wherever possible
6. 🔢 SPECIFIC VALUES: Use exact numbers with units ($, %, GB, days, etc.) - NO vague terms
7. 📋 SCANNABLE FORMAT: H3 subheadings (###) for each distinct aspect/category
8. 🚫 NO FLUFF: No introductions, conclusions, summaries, or transitional phrases
9. ✅ VERIFIED FACTS: Every data point must come directly from research - NO assumptions
10. 💡 ACTIONABLE: Include practical implications/applications of each data point

CONTENT HIERARCHY (strictly follow):
### Data Category 1
- Specific data point: [number/stat/fact]
- Related metric: [number/stat/fact]
- Key comparison: [number vs number]
- Practical insight: [actionable takeaway]

### Data Category 2
1. Numbered step/item: [specific data]
2. Next step/item: [specific data]
3. Next step/item: [specific data]

### Data Category 3
- Feature/aspect: [specifications + numbers]
- Cost/pricing: [exact amounts + comparisons]
- Performance: [metrics + benchmarks]

FORMAT RULES:
- Maximum 2-sentence explanation per category (only if absolutely necessary for context)
- Every bullet must contain at least one specific number, date, or quantifiable fact
- Use bold (**text**) only for key metrics and numbers
- Organize by importance/relevance to user intent
- Include data source indicators when available (e.g., "2024 study:", "Q4 2023:")

FAIL CONDITIONS (return "INSUFFICIENT_DATA" if):
- Research lacks 5+ specific data points
- No concrete numbers, statistics, or metrics available
- Information is too general or vague
- Cannot create meaningful comparisons or structured data

Write 400-600 words of PURE DATA CONTENT NOW (markdown format):
"""

    messages = [{"role": "user", "content": prompt}]
    # FIXED: Increased max_tokens from 1500 to 4096
    response, error = call_grok(messages, max_tokens=4096, temperature=0.5)
    
    if error:
        return None, error
    
    # ENHANCED: Robust check for insufficient data marker
    if response and ("INSUFFICIENT_DATA" in response or len(response.strip()) < 150):
        return None, "INSUFFICIENT_DATA: Unable to generate meaningful content from available research"
    
    return response, None

def generate_table_content(section, research_context, user_intent):
    """Generate comprehensive comparison tables with maximum data extraction"""
    if not grok_key or not section.get('needs_table'):
        return None, None
    
    # Get research data for this section
    section_research = ""
    if 'query_id' in section and section['query_id'] in research_context:
        research_data = research_context[section['query_id']]
        if 'result' in research_data:
            result = research_data['result']
            if isinstance(result, dict) and 'choices' in result:
                section_research = result['choices'][0]['message']['content']
            elif isinstance(result, str):
                section_research = result
    
    # ENHANCED: Stricter check for table-worthy data
    if not section_research or len(section_research.strip()) < 100:
        return None, "INSUFFICIENT_DATA: Not enough data for table generation"
    
    # Check for actual structured data potential
    if "error" in section_research.lower():
        return None, "INSUFFICIENT_DATA: Research contains errors"
    
    # ENHANCED: Aggressive table-focused prompt with user intent
    prompt = f"""Create an EXHAUSTIVE comparison table extracting MAXIMUM data from research.

SECTION: {section['title']}

USER INTENT (Optimize table structure for this):
{user_intent}

RESEARCH DATA:
{section_research}

CRITICAL TABLE REQUIREMENTS:
1. 🎯 MAXIMUM DATA EXTRACTION: Include EVERY extractable data point, metric, feature, price, specification
2. 📊 RICH STRUCTURE: Create 5-10 columns (more columns = better coverage)
3. 📈 COMPREHENSIVE ROWS: Include 10-20 rows with ACTUAL specific data (more rows = better)
4. ✅ NO PLACEHOLDERS: Use only concrete values - NEVER use "Varies", "N/A", "TBD", "Contact", "Coming soon"
5. ❌ INSUFFICIENT DATA CHECK: If you cannot create a table with at least 5 columns and 8 rows of REAL data, return "INSUFFICIENT_DATA"
6. 🔢 UNITS REQUIRED: Include units with all numbers ($, %, GB, MB, days, hours, etc.)
7. 📋 LOGICAL SORTING: Sort rows meaningfully (price low-to-high, popularity, alphabetically, chronologically)
8. 💰 PRICING PRIORITY: Always include pricing/cost data if available
9. 📊 COMPARATIVE METRICS: Include quantitative comparisons wherever possible
10. 🎨 ACTIONABLE STRUCTURE: Table should enable direct comparison and decision-making

OPTIMAL TABLE TYPES BY USER INTENT:
- Comparison: Product/Service A vs B vs C vs D (features, specs, prices)
- Pricing: Tiers/Plans with exact prices, limits, features included/excluded
- Statistical: Time periods with metrics, growth rates, market share
- Specifications: Models/versions with technical specs, performance, compatibility
- Process: Steps with time required, costs, resources, tools needed
- Ranking: Top items ranked with scores, ratings, key differentiators

EXAMPLE HIGH-QUALITY TABLE STRUCTURE:
{{
    "table_title": "Enterprise Cloud Storage Solutions Comparison 2024",
    "headers": ["Provider", "Free Tier", "Business Plan", "Enterprise Plan", "Storage Limit", "File Size Limit", "Security Features", "Integration Options", "Support Level"],
    "rows": [
        ["Google Workspace", "15 GB", "$12/user/mo", "$18/user/mo", "Unlimited", "5 TB", "2FA, DLP, Vault", "1000+ apps", "24/7 Priority"],
        ["Microsoft 365", "5 GB", "$12.50/user/mo", "$22/user/mo", "Unlimited", "250 GB", "MFA, ATP, eDiscovery", "Office Suite", "24/7 Priority"],
        ["Dropbox Business", "2 GB", "$15/user/mo", "$24/user/mo", "Unlimited", "50 GB", "SSO, Device Approval", "Zapier, Slack", "24/5 Standard"],
        ["Box Enterprise", "10 GB", "$17/user/mo", "Custom", "Unlimited", "150 GB", "Shield, Governance", "1500+ apps", "24/7 Dedicated"]
    ]
}}

STRICT JSON FORMAT (return ONLY this):
{{
    "table_title": "Highly Specific, Descriptive Table Title with Year/Context",
    "headers": ["Column 1", "Column 2", "Column 3", "Column 4", "Column 5", "Column 6", "Column 7"],
    "rows": [
        ["Specific value 1", "Specific value 2", "Specific value 3", "Specific value 4", "Specific value 5", "Specific value 6", "Specific value 7"],
        ["Actual data row 2 with real values"],
        ["Actual data row 3 with real values"],
        (minimum 8 rows, target 10-15 rows with complete data)
    ]
}}

FAIL CONDITIONS (return "INSUFFICIENT_DATA" if):
- Cannot create at least 5 columns with distinct data types
- Cannot populate at least 8 rows with specific, non-placeholder values
- Data is too sparse or inconsistent for meaningful comparison
- Missing critical comparative dimensions

Generate MAXIMUM-DATA table NOW (return ONLY valid JSON, extract ALL available information):
"""

    messages = [{"role": "user", "content": prompt}]
    # FIXED: Increased max_tokens from 3000 to 8192 for larger tables
    response, error = call_grok(messages, max_tokens=8192, temperature=0.3)
    
    if error:
        return None, error
    
    # ENHANCED: Check for insufficient data marker
    if response and "INSUFFICIENT_DATA" in response:
        return None, "INSUFFICIENT_DATA: Not enough data for comprehensive table"
    
    try:
        # Extract JSON
        json_match = response
        if "```json" in response:
            json_match = response.split("```json")[1].split("```")[0].strip()
        elif "```" in response:
            json_match = response.split("```")[1].split("```")[0].strip()
        
        table_data = json.loads(json_match)
        
        # ENHANCED: Validate table has substantial data
        if not table_data.get('rows') or len(table_data['rows']) < 5:
            return None, "INSUFFICIENT_DATA: Table has too few data rows"
        
        if not table_data.get('headers') or len(table_data['headers']) < 4:
            return None, "INSUFFICIENT_DATA: Table has too few columns"
        
        # Check for placeholder values
        rows = table_data.get('rows', [])
        for row in rows[:3]:  # Check first 3 rows
            for cell in row:
                cell_str = str(cell).lower()
                if any(placeholder in cell_str for placeholder in ['varies', 'n/a', 'tbd', 'contact', 'coming soon', 'see website']):
                    return None, "INSUFFICIENT_DATA: Table contains placeholder values"
        
        return table_data, None
    except Exception as e:
        return None, f"Table generation error: {str(e)}"

def clean_markdown_symbols(text):
    """Comprehensive function to remove all markdown symbols from text"""
    if not text:
        return ""
    
    # Convert to string if not already
    text = str(text)
    
    # Remove all markdown symbols in specific order to avoid partial matches
    text = text.replace('###', '')
    text = text.replace('##', '')
    text = text.replace('#', '')
    text = text.replace('**', '')
    text = text.replace('__', '')
    text = text.replace('*', '')
    text = text.replace('_', '')
    
    return text.strip()

def export_to_docx(title, meta_desc, sections, keywords):
    """Export article to properly formatted DOCX (no markdown symbols)"""
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = clean_markdown_symbols(title)
    core_properties.keywords = ", ".join([clean_markdown_symbols(k) for k in keywords]) if keywords else ""
    
    # Title (H1) - Only one H1 in the document - CLEANED
    h1 = doc.add_heading(clean_markdown_symbols(title), level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Meta description - CLEANED
    if meta_desc:
        meta = doc.add_paragraph(clean_markdown_symbols(meta_desc))
        meta.italic = True
        meta_format = meta.paragraph_format
        meta_format.space_after = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Add sections
    for section_data in sections:
        section = section_data['section']
        content = section_data['content']
        
        # Section heading (H2) - FULLY CLEANED
        section_title_clean = clean_markdown_symbols(section['title'])
        h2 = doc.add_heading(section_title_clean, level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Process content line by line
        content_lines = content.split('\n')
        
        for line in content_lines:
            line = line.strip()
            if not line:
                continue
            
            # H3 subheadings - CLEANED
            if line.startswith('###'):
                subheading_text = clean_markdown_symbols(line.replace('###', ''))
                if subheading_text:
                    h3 = doc.add_heading(subheading_text, level=3)
                    h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Bullet points - CLEANED
            elif line.startswith('- ') or line.startswith('* '):
                text = clean_markdown_symbols(line[2:].strip())
                if text:
                    doc.add_paragraph(text, style='List Bullet')
            # Numbered lists - CLEANED
            elif line[0].isdigit() and '. ' in line:
                text = line.split('. ', 1)[1] if '. ' in line else line
                text = clean_markdown_symbols(text)
                if text:
                    doc.add_paragraph(text, style='List Number')
            # Regular paragraph - FULLY CLEANED
            else:
                text = clean_markdown_symbols(line)
                if text:
                    doc.add_paragraph(text)
        
        # Add table if available
        if 'table' in section_data:
            table_data = section_data['table']
            doc.add_paragraph()  # Spacing
            
            # Table title (H3) - CLEANED
            table_title_clean = clean_markdown_symbols(table_data.get('table_title', 'Data Table'))
            table_title = doc.add_heading(table_title_clean, level=3)
            table_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Create table
            headers = table_data['headers']
            rows = table_data['rows']
            
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Light Grid Accent 1'
            
            # Header row - FULLY CLEANED
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                clean_header = clean_markdown_symbols(header)
                header_cells[i].text = clean_header
                # Bold header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows - FULLY CLEANED
            for row_idx, row_data in enumerate(rows):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, cell_value in enumerate(row_data):
                    clean_value = clean_markdown_symbols(cell_value)
                    row_cells[col_idx].text = clean_value
        
        # Infographic note - CLEANED
        if section.get('needs_infographic'):
            doc.add_paragraph()
            infographic_desc = clean_markdown_symbols(section.get('infographic_description', 'Visual recommended'))
            info_note = doc.add_paragraph(f"💡 Infographic: {infographic_desc}")
            info_note.italic = True
        
        doc.add_paragraph()  # Section spacing
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_to_html(title, meta_desc, sections, keywords):
    """Export article to clean HTML (no html/head tags, one H1, no markdown symbols)"""
    html = f"<article>\n"
    
    # Title (H1) - Only one H1 - CLEANED
    html += f"  <h1>{clean_markdown_symbols(title)}</h1>\n"
    
    # Meta description - CLEANED
    if meta_desc:
        html += f"  <p class=\"meta-description\"><em>{clean_markdown_symbols(meta_desc)}</em></p>\n"
    
    html += "\n"
    
    # Add sections
    for section_data in sections:
        section = section_data['section']
        content = section_data['content']
        
        # Section heading (H2) - CLEANED
        html += f"  <h2>{clean_markdown_symbols(section['title'])}</h2>\n"
        
        # Process content line by line and remove markdown formatting
        content_lines = content.split('\n')
        in_list = False
        list_type = None
        
        for line in content_lines:
            line = line.strip()
            if not line:
                if in_list:
                    html += f"  </{list_type}>\n"
                    in_list = False
                continue
            
            # H3 subheadings - FULLY CLEANED
            if line.startswith('###'):
                if in_list:
                    html += f"  </{list_type}>\n"
                    in_list = False
                subheading_text = clean_markdown_symbols(line.replace('###', ''))
                if subheading_text:
                    html += f"  <h3>{subheading_text}</h3>\n"
            # Bullet points - FULLY CLEANED
            elif line.startswith('- ') or line.startswith('* '):
                text = clean_markdown_symbols(line[2:].strip())
                if not in_list or list_type != 'ul':
                    if in_list:
                        html += f"  </{list_type}>\n"
                    html += "  <ul>\n"
                    in_list = True
                    list_type = 'ul'
                if text:
                    html += f"    <li>{text}</li>\n"
            # Numbered lists - FULLY CLEANED
            elif line[0].isdigit() and '. ' in line:
                text = line.split('. ', 1)[1] if '. ' in line else line
                text = clean_markdown_symbols(text)
                if not in_list or list_type != 'ol':
                    if in_list:
                        html += f"  </{list_type}>\n"
                    html += "  <ol>\n"
                    in_list = True
                    list_type = 'ol'
                if text:
                    html += f"    <li>{text}</li>\n"
            # Regular paragraph - FULLY CLEANED
            else:
                if in_list:
                    html += f"  </{list_type}>\n"
                    in_list = False
                text = clean_markdown_symbols(line)
                if text:
                    html += f"  <p>{text}</p>\n"
        
        if in_list:
            html += f"  </{list_type}>\n"
        
        # Add table if available
        if 'table' in section_data:
            table_data = section_data['table']
            html += "\n"
            
            # Table title (H3) - FULLY CLEANED
            table_title = clean_markdown_symbols(table_data.get('table_title', 'Data Table'))
            html += f"  <h3>{table_title}</h3>\n"
            
            # Create table
            html += "  <table>\n"
            html += "    <thead>\n      <tr>\n"
            for header in table_data['headers']:
                clean_header = clean_markdown_symbols(header)
                html += f"        <th>{clean_header}</th>\n"
            html += "      </tr>\n    </thead>\n"
            
            html += "    <tbody>\n"
            for row in table_data['rows']:
                html += "      <tr>\n"
                for cell in row:
                    clean_cell = clean_markdown_symbols(cell)
                    html += f"        <td>{clean_cell}</td>\n"
                html += "      </tr>\n"
            html += "    </tbody>\n"
            html += "  </table>\n"
        
        # Infographic note - CLEANED
        if section.get('needs_infographic'):
            infographic_desc = clean_markdown_symbols(section.get('infographic_description', 'Visual recommended'))
            html += f"  <p class=\"infographic-note\"><em>💡 Infographic: {infographic_desc}</em></p>\n"
        
        html += "\n"
    
    html += "</article>"
    return html

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF: {e}"

def generate_query_fanout(topic):
    """Generate diverse research queries"""
    if not model:
        return None, "Gemini API key required"
    
    prompt = f"""Generate 15-20 diverse, specific research queries for comprehensive coverage of: "{topic}"

REQUIREMENTS:
1. Mix of query types:
   - Definitional ("What is...", "How does...")
   - Comparative ("X vs Y", "Compare...")
   - Practical ("How to...", "Steps to...")
   - Statistical ("Statistics on...", "Data about...")
   - Historical ("History of...", "Evolution of...")
   - Future-looking ("Trends in...", "Future of...")
   - Problem-solving ("Challenges in...", "Solutions for...")

2. Vary specificity (broad to narrow)
3. Target different user intents
4. Include industry-specific queries
5. Add location/region-specific queries if relevant

Return ONLY a JSON array of query strings:
["query1", "query2", "query3", ...]

Generate queries NOW:
"""

    try:
        response = model.generate_content(prompt)
        result = response.text.strip()
        
        # Extract JSON
        if "```json" in result:
            result = result.split("```json")[1].split("```")[0].strip()
        elif "```" in result:
            result = result.split("```")[1].split("```")[0].strip()
        
        queries = json.loads(result)
        return queries, None
    except Exception as e:
        return None, f"Query generation error: {str(e)}"

def analyze_pdf_content(pdf_text):
    """Analyze PDF and extract key topics"""
    if not model:
        return None, "Gemini API key required"
    
    prompt = f"""Analyze this PDF content and extract 5-10 key topics/themes for research:

PDF CONTENT:
{pdf_text[:4000]}  # Limit context

Return ONLY a JSON array of topic strings:
["topic1", "topic2", "topic3", ...]

Topics should be:
- Specific and focused
- Researchable
- Actionable for content creation

Extract topics NOW:
"""

    try:
        response = model.generate_content(prompt)
        result = response.text.strip()
        
        # Extract JSON
        if "```json" in result:
            result = result.split("```json")[1].split("```")[0].strip()
        elif "```" in result:
            result = result.split("```")[1].split("```")[0].strip()
        
        topics = json.loads(result)
        return topics, None
    except Exception as e:
        return None, f"PDF analysis error: {str(e)}"

# Main App Layout
tab1, tab2, tab3, tab4 = st.tabs(["🔍 Query Fan-Out", "📄 PDF Analysis", "✅ Fact Checker", "📊 Research Dashboard"])

# TAB 1: Query Fan-Out
with tab1:
    st.header("Query Fan-Out Generator")
    
    query_input = st.text_area(
        "Enter your research topic or query:",
        height=100,
        placeholder="E.g., 'AI in healthcare', 'Sustainable energy solutions', 'Remote work trends 2024'"
    )
    
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("🚀 Generate Queries", use_container_width=True):
            if not query_input:
                st.error("Please enter a topic")
            elif not model:
                st.error("Please provide Gemini API key")
            else:
                with st.spinner("Generating diverse research queries..."):
                    queries, error = generate_query_fanout(query_input)
                    
                    if error:
                        st.error(f"Error: {error}")
                    else:
                        st.session_state.fanout_results = queries
                        st.session_state.generation_details = {
                            'topic': query_input,
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            'model': gemini_model
                        }
                        st.session_state.user_query = query_input
                        st.success(f"✅ Generated {len(queries)} queries!")
                        st.rerun()
    
    with col2:
        if st.session_state.fanout_results:
            if st.button("🗑️ Clear Results", use_container_width=True):
                st.session_state.fanout_results = None
                st.session_state.generation_details = None
                st.session_state.research_results = {}
                st.session_state.selected_queries = set()
                st.session_state.user_query = ''
                st.rerun()
    
    # Display Results
    if st.session_state.fanout_results:
        st.divider()
        
        # Generation info
        if st.session_state.generation_details:
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Topic", st.session_state.generation_details.get('topic', 'N/A'))
            with col2:
                st.metric("Queries Generated", len(st.session_state.fanout_results))
            with col3:
                st.metric("Model Used", st.session_state.generation_details.get('model', 'N/A'))
        
        st.divider()
        st.subheader("📋 Generated Queries")
        
        # Query selection
        st.info("💡 Select queries below to perform research")
        
        # Select all / Deselect all buttons
        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Select All"):
                st.session_state.selected_queries = set(range(len(st.session_state.fanout_results)))
                st.rerun()
        with col2:
            if st.button("❌ Deselect All"):
                st.session_state.selected_queries = set()
                st.rerun()
        
        # Display queries with checkboxes
        for idx, query in enumerate(st.session_state.fanout_results):
            is_selected = idx in st.session_state.selected_queries
            
            col1, col2 = st.columns([0.1, 0.9])
            with col1:
                if st.checkbox("", value=is_selected, key=f"query_{idx}", label_visibility="collapsed"):
                    st.session_state.selected_queries.add(idx)
                else:
                    st.session_state.selected_queries.discard(idx)
            with col2:
                st.markdown(f"**{idx + 1}.** {query}")
        
        # Research button
        if st.session_state.selected_queries:
            st.divider()
            st.info(f"🎯 {len(st.session_state.selected_queries)} queries selected for research")
            
            if st.button("🔬 Start Research on Selected Queries", use_container_width=True):
                if not perplexity_key:
                    st.error("❌ Perplexity API key required for research")
                else:
                    # Clear previous research results
                    st.session_state.research_results = {}
                    
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    selected_list = sorted(list(st.session_state.selected_queries))
                    
                    for i, query_idx in enumerate(selected_list):
                        query = st.session_state.fanout_results[query_idx]
                        status_text.text(f"Researching: {query}")
                        
                        # Call Perplexity
                        result = call_perplexity(query)
                        
                        # Store result
                        st.session_state.research_results[query_idx] = {
                            'query': query,
                            'result': result,
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
                        
                        # Update progress
                        progress = (i + 1) / len(selected_list)
                        progress_bar.progress(progress)
                        
                        time.sleep(0.5)  # Rate limiting
                    
                    status_text.text("✅ Research complete!")
                    progress_bar.progress(1.0)
                    st.success(f"🎉 Completed research on {len(selected_list)} queries!")
                    time.sleep(1)
                    st.rerun()
        
        # FIXED: Display Research Results Section - Now displays correctly after research
        if st.session_state.research_results:
            st.divider()
            st.subheader("🔍 Research Results")
            
            st.info(f"✅ Completed {len(st.session_state.research_results)} research queries")
            
            # Show research results in expandable sections
            for query_id in sorted(st.session_state.research_results.keys()):
                research_data = st.session_state.research_results[query_id]
                with st.expander(f"📊 {research_data['query']}", expanded=False):
                    result = research_data['result']
                    
                    # Display result
                    if isinstance(result, dict):
                        if 'error' in result:
                            st.error(f"❌ {result['error']}")
                        elif 'choices' in result:
                            content = result['choices'][0]['message']['content']
                            st.markdown(content)
                            
                            # Show timestamp
                            st.caption(f"📅 Researched: {research_data['timestamp']}")
                        else:
                            st.json(result)
                    else:
                        st.markdown(result)
                    
                    st.caption(f"🔑 Query ID: {query_id}")
        
        # Export options
        if st.session_state.fanout_results:
            st.divider()
            st.subheader("📥 Export Options")
            
            # Prepare export data - safe access to avoid KeyError
            if st.session_state.generation_details:
                export_data = {
                    'topic': st.session_state.generation_details['topic'],
                    'timestamp': st.session_state.generation_details['timestamp'],
                    'model': st.session_state.generation_details['model'],
                    'queries': st.session_state.fanout_results
                }
            else:
                export_data = {
                    'topic': 'Unknown Topic',
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'model': 'Unknown',
                    'queries': st.session_state.fanout_results
                }
            
            col1, col2 = st.columns(2)
            
            with col1:
                # JSON export
                json_data = json.dumps(export_data, indent=2)
                st.download_button(
                    "📄 Download JSON",
                    data=json_data,
                    file_name=f"queries_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    use_container_width=True
                )
            
            with col2:
                # CSV export
                df = pd.DataFrame({
                    'Query_Number': range(1, len(st.session_state.fanout_results) + 1),
                    'Query': st.session_state.fanout_results
                })
                csv_data = df.to_csv(index=False)
                st.download_button(
                    "📊 Download CSV",
                    data=csv_data,
                    file_name=f"queries_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )

# TAB 2: PDF Analysis
with tab2:
    st.header("PDF Content Analysis")
    
    uploaded_pdf = st.file_uploader("Upload PDF Document", type=['pdf'])
    
    if uploaded_pdf:
        with st.spinner("Extracting PDF content..."):
            pdf_text = extract_text_from_pdf(uploaded_pdf)
            
            if "Error" in pdf_text:
                st.error(pdf_text)
            else:
                st.success(f"✅ Extracted {len(pdf_text)} characters")
                
                with st.expander("📄 View Extracted Text"):
                    st.text_area("PDF Content", pdf_text[:2000], height=300, disabled=True)
                
                if st.button("🔍 Analyze Topics"):
                    if not model:
                        st.error("Gemini API key required")
                    else:
                        with st.spinner("Analyzing PDF content..."):
                            topics, error = analyze_pdf_content(pdf_text)
                            
                            if error:
                                st.error(f"Error: {error}")
                            else:
                                st.session_state.pdf_analysis = {
                                    'topics': topics,
                                    'pdf_name': uploaded_pdf.name,
                                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                }
                                st.success(f"✅ Identified {len(topics)} key topics!")
                                st.rerun()
        
        # Display analyzed topics
        if st.session_state.pdf_analysis:
            st.divider()
            st.subheader("📊 Identified Topics")
            
            topics = st.session_state.pdf_analysis['topics']
            
            for idx, topic in enumerate(topics):
                col1, col2 = st.columns([0.8, 0.2])
                with col1:
                    st.markdown(f"**{idx + 1}.** {topic}")
                with col2:
                    if st.button("🔬 Research", key=f"research_topic_{idx}"):
                        # Add to enhanced topics for research
                        if topic not in st.session_state.enhanced_topics:
                            st.session_state.enhanced_topics.append(topic)
                            st.success(f"Added '{topic}' to research queue")

# TAB 3: Fact Checker
with tab3:
    st.header("Fact Checker")
    
    claim_input = st.text_area(
        "Enter claim to verify:",
        height=100,
        placeholder="E.g., 'AI will replace 50% of jobs by 2030'"
    )
    
    if st.button("✅ Verify Claim", use_container_width=True):
        if not claim_input:
            st.error("Please enter a claim")
        elif not perplexity_key:
            st.error("Perplexity API key required")
        else:
            with st.spinner("Verifying claim..."):
                system_prompt = "Fact-check this claim. Provide evidence-based verification with sources. Rate confidence as: HIGH, MEDIUM, or LOW. Format: VERDICT | CONFIDENCE | EXPLANATION"
                result = call_perplexity(claim_input, system_prompt)
                
                if 'error' in result:
                    st.error(f"Error: {result['error']}")
                elif 'choices' in result:
                    content = result['choices'][0]['message']['content']
                    st.markdown("### Verification Result")
                    st.markdown(content)
                else:
                    st.json(result)

# TAB 4: Research Dashboard & Content Generation
with tab4:
    st.header("📊 Research Dashboard & AI Content Generator")
    
    # Check if research results exist
    if not st.session_state.research_results:
        st.warning("⚠️ No research data available. Please perform research in the 'Query Fan-Out' tab first.")
        st.info("💡 Go to 'Query Fan-Out' tab → Generate queries → Select queries → Start Research")
    else:
        # Show research summary
        st.success(f"✅ {len(st.session_state.research_results)} research queries available for content generation")
        
        with st.expander("📊 View Available Research Data"):
            for query_id, research_data in st.session_state.research_results.items():
                st.markdown(f"**{query_id + 1}.** {research_data['query']}")
        
        st.divider()
        st.subheader("📝 AI Content Generator")
        
        # Content generation inputs
        with st.container():
            content_topic = st.text_input(
                "Article Title (H1):",
                value=st.session_state.user_query if st.session_state.user_query else "",
                placeholder="E.g., 'Complete Guide to AI in Healthcare 2024'"
            )
            
            # ENHANCED: User Intent input with clear emphasis on its importance
            user_intent = st.text_area(
                "User Intent / Purpose (CRITICAL - Determines output format):",
                height=120,
                placeholder="E.g., 'Compare pricing and features of different tools with detailed tables', 'Show step-by-step implementation with specific technical data', 'Provide statistical comparison of market leaders with quantitative metrics', 'Create data-heavy analysis with minimal narrative'",
                help="⚠️ IMPORTANT: This field determines how content is generated. Be specific about wanting tables, data points, comparisons, etc. The more detail you provide about desired output format, the better the results."
            )
            
            if st.button("📝 Generate Article Structure", use_container_width=True):
                if not content_topic:
                    st.error("❌ Please provide an article title")
                elif not user_intent:
                    st.error("❌ Please describe the user intent/purpose - this is critical for generating appropriate content format")
                elif not grok_key:
                    st.error("❌ Grok API key required")
                else:
                    with st.spinner("Creating article structure from your research queries..."):
                        structure, error = generate_content_structure_from_queries(
                            st.session_state.research_results,
                            content_topic,
                            st.session_state.selected_queries
                        )
                        
                        if error:
                            st.error(f"❌ Structure generation failed: {error}")
                        else:
                            st.session_state.content_structure = structure
                            st.session_state.user_intent = user_intent  # Store intent
                            st.success("✅ Article structure created from your queries!")
                            st.rerun()
            
            # Display and Edit Structure
            if st.session_state.content_structure:
                st.divider()
                st.subheader("📋 Article Structure (Based on Your Queries)")
                
                structure = st.session_state.content_structure
                
                # Display title and meta
                st.markdown(f"**Title (H1):** {structure['article_title']}")
                st.markdown(f"**Meta Description:** {structure['meta_description']}")
                
                st.divider()
                st.markdown("**Sections (H2 Headings - Exact Query Text):**")
                
                # Editable sections
                edited_sections = []
                
                for idx, section in enumerate(structure['sections']):
                    with st.expander(f"Section {idx + 1}: {section['title']}", expanded=True):
                        col1, col2 = st.columns([3, 1])
                        
                        with col1:
                            section_title = st.text_input(
                                "Section Title (from query)",
                                value=section['title'],
                                key=f"section_title_{idx}",
                                label_visibility="collapsed"
                            )
                        
                        with col2:
                            needs_table = st.checkbox(
                                "Include Table",
                                value=section.get('needs_table', True),
                                key=f"table_{idx}"
                            )
                        
                        section_copy = section.copy()
                        section_copy['title'] = section_title
                        section_copy['needs_table'] = needs_table
                        edited_sections.append(section_copy)
                
                structure['sections'] = edited_sections
                st.session_state.content_structure = structure
                
                # Generate Content Button
                st.divider()
                
                if st.button("🚀 Generate Full Article", use_container_width=True, type="primary"):
                    if not grok_key:
                        st.error("❌ Grok API key required")
                    else:
                        # Clear previous content
                        st.session_state.generated_content = {}
                        
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        # FIXED: Get user intent from session state (guaranteed to exist now)
                        user_intent = st.session_state.get('user_intent', 'Provide comprehensive data-driven information')
                        
                        successful_sections = 0
                        skipped_sections = []
                        
                        for idx, section in enumerate(edited_sections):
                            section_key = f"section_{idx}"
                            status_text.text(f"Generating: {section['title']}")
                            
                            # ENHANCED: Generate section content with user intent
                            content, error = generate_section_content(
                                section,
                                st.session_state.research_results,
                                content_topic,
                                user_intent
                            )
                            
                            # ENHANCED: Skip section if insufficient data
                            if error and "INSUFFICIENT_DATA" in error:
                                status_text.text(f"⚠️ Skipping '{section['title']}' - insufficient research data")
                                skipped_sections.append(section['title'])
                                time.sleep(1)
                                continue
                            
                            if error:
                                st.warning(f"⚠️ Error generating '{section['title']}': {error}")
                                time.sleep(1)
                                continue
                            
                            section_data = {
                                'content': content,
                                'section': section
                            }
                            
                            # ENHANCED: Generate table if needed with user intent
                            if section.get('needs_table'):
                                table, table_error = generate_table_content(
                                    section,
                                    st.session_state.research_results,
                                    user_intent
                                )
                                if table and not table_error:
                                    section_data['table'] = table
                                elif table_error and "INSUFFICIENT_DATA" in table_error:
                                    status_text.text(f"ℹ️ Table skipped for '{section['title']}' - insufficient data")
                                elif table_error:
                                    st.info(f"ℹ️ Could not generate table for '{section['title']}': {table_error}")
                            
                            st.session_state.generated_content[section_key] = section_data
                            successful_sections += 1
                            
                            # Update progress
                            progress = (idx + 1) / len(edited_sections)
                            progress_bar.progress(progress)
                            
                            time.sleep(1)  # Rate limiting
                        
                        status_text.text("✅ Article generation complete!")
                        progress_bar.progress(1.0)
                        
                        if successful_sections > 0:
                            st.success(f"🎉 Successfully generated {successful_sections} sections!")
                            if skipped_sections:
                                st.info(f"ℹ️ Skipped {len(skipped_sections)} sections due to insufficient data: {', '.join(skipped_sections[:3])}{'...' if len(skipped_sections) > 3 else ''}")
                        else:
                            st.error("❌ No sections could be generated. The research data may be insufficient.")
                        
                        time.sleep(2)
                        st.rerun()
            
            # Display Generated Content
            if st.session_state.generated_content:
                st.divider()
                st.subheader("📄 Generated Article")
                
                # Article preview - safe access to structure
                article_title = st.session_state.content_structure.get('article_title', 'Untitled Article')
                meta_desc = st.session_state.content_structure.get('meta_description', '')
                
                st.markdown(f"# {article_title}")
                if meta_desc:
                    st.caption(f"*{meta_desc}*")
                st.divider()
                
                total_words = 0
                
                for idx, section_key in enumerate(sorted(st.session_state.generated_content.keys())):
                    section_data = st.session_state.generated_content.get(section_key, {})
                    if not section_data:
                        continue
                        
                    section = section_data.get('section', {})
                    content = section_data.get('content', '')
                    
                    section_title = section.get('title', f'Section {idx + 1}')
                    st.markdown(f"## {section_title}")
                    st.markdown(content)
                    
                    # Show table if generated
                    if 'table' in section_data:
                        table = section_data.get('table', {})
                        table_title = table.get('table_title', 'Comparison Table')
                        st.markdown(f"### {table_title}")
                        
                        # Create dataframe for display
                        if 'headers' in table and 'rows' in table:
                            df = pd.DataFrame(table['rows'], columns=table['headers'])
                            st.dataframe(df, use_container_width=True, hide_index=True)
                    
                    # Show infographic suggestion
                    if section.get('needs_infographic'):
                        infographic_desc = section.get('infographic_description', 'Visual representation recommended')
                        st.info(f"💡 **Infographic Suggestion:** {infographic_desc}")
                    
                    st.divider()
                    
                    total_words += len(content.split())
                
                # Article stats
                st.success(f"📊 **Total Word Count:** {total_words:,} words")
                
                # Export options
                st.subheader("📥 Export Article")
                
                # Prepare data for export
                sections_list = []
                for section_key in sorted(st.session_state.generated_content.keys()):
                    section_data = st.session_state.generated_content.get(section_key)
                    if section_data:
                        sections_list.append(section_data)
                
                # Get keywords
                article_keywords = st.session_state.content_structure.get('semantic_keywords', [])
                
                # Export buttons in columns
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    # FIXED: DOCX Export with complete markdown cleaning
                    if DOCX_AVAILABLE:
                        docx_buffer = export_to_docx(
                            st.session_state.content_structure.get('article_title', 'Article'),
                            st.session_state.content_structure.get('meta_description', ''),
                            sections_list,
                            article_keywords
                        )
                        
                        if docx_buffer:
                            st.download_button(
                                "📄 Download DOCX",
                                data=docx_buffer,
                                file_name=f"{content_topic.replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    else:
                        st.warning("Install python-docx for DOCX export")
                
                with col2:
                    # FIXED: Clean HTML Export with complete markdown cleaning
                    html_content = export_to_html(
                        st.session_state.content_structure.get('article_title', 'Article'),
                        st.session_state.content_structure.get('meta_description', ''),
                        sections_list,
                        article_keywords
                    )
                    
                    st.download_button(
                        "🌐 Download HTML",
                        data=html_content.encode('utf-8'),
                        file_name=f"{content_topic.replace(' ', '_')}.html",
                        mime="text/html",
                        use_container_width=True
                    )
                
                with col3:
                    # FIXED: Text Export with complete markdown cleaning
                    article_title = st.session_state.content_structure.get('article_title', 'Article')
                    article_meta = st.session_state.content_structure.get('meta_description', '')
                    
                    full_article = f"{clean_markdown_symbols(article_title)}\n\n"
                    if article_meta:
                        full_article += f"{clean_markdown_symbols(article_meta)}\n\n"
                    full_article += "---\n\n"
                    
                    for section_key in sorted(st.session_state.generated_content.keys()):
                        section_data = st.session_state.generated_content.get(section_key, {})
                        if not section_data:
                            continue
                            
                        section = section_data.get('section', {})
                        
                        # Clean section title completely
                        section_title = clean_markdown_symbols(section.get('title', 'Section'))
                        full_article += f"{section_title}\n\n"
                        
                        # Clean content completely - remove ALL markdown symbols
                        content_clean = clean_markdown_symbols(section_data.get('content', ''))
                        full_article += f"{content_clean}\n\n"
                        
                        if 'table' in section_data:
                            table = section_data.get('table', {})
                            table_title = clean_markdown_symbols(table.get('table_title', 'Table'))
                            full_article += f"{table_title}\n\n"
                            
                            headers = table.get('headers', [])
                            rows = table.get('rows', [])
                            
                            if headers:
                                headers_clean = [clean_markdown_symbols(h) for h in headers]
                                full_article += " | ".join(headers_clean) + "\n"
                                full_article += " | ".join(["---"] * len(headers_clean)) + "\n"
                                
                                for row in rows:
                                    clean_row = [clean_markdown_symbols(cell) for cell in row]
                                    full_article += " | ".join(clean_row) + "\n"
                                full_article += "\n"
                        
                        if section.get('needs_infographic'):
                            infographic = clean_markdown_symbols(section.get('infographic_description', 'Visual recommended'))
                            full_article += f"💡 Infographic: {infographic}\n\n"
                        
                        full_article += "---\n\n"
                    
                    st.download_button(
                        "📝 Download Text",
                        data=full_article.encode('utf-8'),
                        file_name=f"{content_topic.replace(' ', '_')}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                
                # Show format info
                with st.expander("ℹ️ Export Format Details"):
                    st.markdown("""
                    **DOCX (Word):**
                    - Full formatting with proper headings
                    - One H1 title, H2 for sections, H3 for subsections
                    - Tables properly formatted with bold headers
                    - ✅ **Clean output with NO markdown symbols (*, #, _)**
                    - Editable in Microsoft Word/Google Docs
                    
                    **HTML (Clean):**
                    - No html, head, or meta tags (ready for CMS)
                    - One H1 title, H2 for sections, H3 for subsections
                    - Semantic HTML structure with proper lists
                    - Tables with thead/tbody markup
                    - ✅ **Clean output with NO markdown symbols (*, #, _)**
                    
                    **Text:**
                    - Plain text format
                    - ✅ **NO formatting symbols or markdown (*, #, _)**
                    - Clean, readable content
                    - Tables in simple text format
                    """)

# Clear all data button
if st.sidebar.button("🗑️ Clear All Data"):
    st.session_state.fanout_results = None
    st.session_state.generation_details = None
    st.session_state.research_results = {}
    st.session_state.selected_queries = set()
    st.session_state.pdf_analysis = None
    st.session_state.enhanced_topics = []
    st.session_state.generated_content = {}
    st.session_state.content_structure = []
    st.session_state.user_query = ''
    st.session_state.user_intent = ''
    st.success("All data cleared!")
    st.rerun()

# Footer
st.markdown("---")
st.markdown("**Qforia Complete Research Platform** - Query Fan-Out, PDF Analysis, Fact Checking, Research Dashboard & AI Content Generation | *Powered by Gemini AI, Perplexity & Grok*")
